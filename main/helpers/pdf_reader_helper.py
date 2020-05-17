import json, datetime
from main.models import File, Document, Site
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from pdf2image import convert_from_path
from main.utils.common_utils import CommonUtils
from main.utils.image_utils import ImageUtils
from main.workers.document_worker import DocumentWorker
from multiprocessing import Process
from docx import Document as WordDocument
from django.db import connection 

class PDFReaderHelper:

    def create_file(self, pdf):

        common_utils = CommonUtils()
        internal_name = common_utils.create_token()+'.pdf'
        
        fs = FileSystemStorage(location=settings.BASE_POOL_PDF) 
        filename = fs.save(internal_name, pdf)

        file_url = fs.url(filename)
        file_db = File()
        file_db.user_name = pdf.name
        file_db.internal_name = internal_name
        file_db.date = datetime.datetime.now()
        file_db.status = file_db.STARTED
        file_db.save()
      
        p = Process(target=self.pdf_process,  args=(file_db.id,))
        p.start()


    def pdf_process(self, file_id):

        connection.close()
        file_db = File.objects.get(id=file_id)

        common_utils = CommonUtils()
        image_utils = ImageUtils()
                
        pages = convert_from_path(settings.BASE_POOL_PDF+"/"+file_db.internal_name, 300)
        file_db.status = file_db.IN_PROCESS
        file_db.save()
        documents_ok = 0

        word_document = WordDocument()

        word_document.add_heading('INFORMACIÓN OBTENIDA POR CEDULA DE CIUDADANIA', 0)
        for page in pages:
            page.save('out.jpg', 'JPEG')
            img = image_utils.process_image_for_ocr('out.jpg')
            fiscal_number, first_name, last_name, sex =  image_utils.get_text_from_document(img)
            
            if fiscal_number != "" and first_name != "" and last_name != "":
                word_document.add_heading('Cedula: '+fiscal_number, level=1)
                word_document.add_heading('Apellidos: '+last_name, level=1)
                word_document.add_heading('Nombres: '+first_name, level=1)
                word_document.add_heading('Sexo: '+sex, level=1)

                document = Document()
                document.file = file_db
                document.fiscal_number = fiscal_number
                document.sex = sex
                document.first_name = first_name
                document.last_name = last_name
                document.save()

                document_worker = DocumentWorker()
                sites = document_worker.run_worker(first_name, last_name)
                
                if len(sites)>0:
                    for site in sites:
                        word_document.add_heading('Pagina: '+site['url'], level=2)
                        word_document.add_heading('Nube de palabras: ', level=2)
                        word_document.add_paragraph(site['glossary'], style='List Bullet')
                        site_db = Site()
                        site_db.document = document
                        site_db.url = site['url']
                        site_db.glossary = site['glossary']
                        site_db.save()
                    documents_ok += 1
                word_document.add_page_break()
                
            else:
                print("Hoja No contiene un cedula colombiana valida")

        if documents_ok > 0:
            file_db.status = file_db.SUCCESS
        else:
            file_db.status = file_db.ERROR
        word_internal_name = common_utils.create_token()+'.docx'   
        word_document.save(settings.BASE_POOL_WORD+"/"+word_internal_name)   
        file_db.report_internal_name = word_internal_name
        file_db.save()  

    def get_files(self):
    
        files = File.objects.all().order_by("-date")
        files_resp = []
        for file in files:
            files_resp.append({
                'id': file.id,
                'date' : str(file.date)[:19],
                'internal_name': file.internal_name,
                'report_internal_name': file.report_internal_name,
                'user_name': file.user_name,
                'status': file.get_status_display()
            })

        return files_resp



